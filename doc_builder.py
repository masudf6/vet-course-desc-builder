from io import BytesIO
from docx import Document
from docx.shared import Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from typing import List, Optional, Union
import re

from models import Unit


def style_table(table, doc, ratios=None):
    """
    Styles `table` with solid borders, cell padding, and
    column widths based on `ratios` of the usable page width.

    - `doc` is the Document() that owns this table.
    - `ratios` is a list of N numbers (one per column); they
       can be arbitrary (e.g. [1,2,1,1,3]) and will be
       normalized to fill 100% of the usable width.
    """
    # 1) Defaults & get usable width
    if ratios is None:
        ratios = [.5, 2, .5, .5, 4]  # tweak to taste
    total = sum(ratios)

    sect = doc.sections[0]
    usable = sect.page_width - sect.left_margin - sect.right_margin

    # 2) Apply “Table Grid” and turn autofit off
    table.style = 'Table Grid'

    # 3) Set column widths proportionally
    for col_idx, (col, ratio) in enumerate(zip(table.columns, ratios)):
        col_width = Emu(int(usable * (ratio/total)))
        for cell in col.cells:
            cell.width = col_width

def build_doc(units: List[Unit], title: str = '', out_path: Optional[str] = None) -> Union[BytesIO, None]:
    doc = Document()

    section = doc.sections[0]
    usable_width = section.page_width - section.left_margin - section.right_margin

    # ── HEADER IMAGE ──
    header = section.header
    hdr_p = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
    run = hdr_p.add_run()
    run.add_picture('header.png', width=usable_width)
    hdr_p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # ── FOOTER ──
    # footer = section.footer

    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    doc.add_heading(f"Course Description: {title}", level=1)

    tbl = doc.add_table(rows=1, cols=5)
    style_table(tbl, doc)  # ← apply our formatting

    # Header row
    hdr = ["Subject Code","Subject Name","Subject Type","Contact Hours","Performance Criteria"]
    for i, txt in enumerate(hdr):
        tbl.rows[0].cells[i].text = txt

    # Data rows
    for u in units:
        cells = tbl.add_row().cells
        cells[0].text = u.code
        cells[1].text = u.name
        cells[2].text = u.type
        cells[3].text = u.hours

        # True bullets instead of plain text
        crit_cell = cells[4]
        # ----- clear out the default/empty paragraphs -----
        for paragraph in list(crit_cell.paragraphs):
            paragraph._element.getparent().remove(paragraph._element)

        cleaned = [re.sub(r'^\d+\.\d+\s*', '', c) for c in u.criteria]
        for c in cleaned:
            p = crit_cell.add_paragraph(style='List Bullet')
            p.add_run(c)

    if out_path:
        doc.save(out_path)
        return None
    else:
        buf = BytesIO()
        doc.save(buf)
        buf.seek(0)
        return buf